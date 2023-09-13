#
# Imports
#
from xml.dom.minidom import Document
import docx
from lorem import *
import function
import ruthinit
from docx2pdf import convert
import aspose.words as aw
from docx.shared import RGBColor

#
# Globals
#
log = ruthinit.log

def build_docx(output_text):
        

    # Create a document
    doc = docx.Document()

    # Add a paragraph to the document
    p = doc.add_paragraph()

    # Add some formatting to the paragraph
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0

    # Add a run to the paragraph
    run = p.add_run("python-docx")

    # Add some formatting to the run
    run.bold = True
    run.italic = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(16)


    # Add more text to the same paragraph
    run = p.add_run(" Tutorial")

    # Format the run
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(16)

    # Add another paragraph (left blank for an empty line)
    doc.add_paragraph()

    # Add another paragraph
    p = doc.add_paragraph()

    # Add a run and format it
    run = p.add_run(output_text)
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)

    # Save the document
    doc.save("output.docx")

def build_word_document(action_items, rca_why, rca_det, inc_timeline):
    log.info("Creating Document")
    #initialize document
    doc = docx.Document()
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "\t\t"
    run = paragraph.add_run()
    run.add_picture("static/ruthlogosmall.png")

    # Main Color: Core Green (75,205,62)
    # Supporting Color 1: Teal (0,151,117)
    # Supporting Color 2: Dark Teal (1,91,126)

    #Main Header
    h0 = doc.add_heading('RUTH Analysis', 0)
    h0.style.font.color.rgb = RGBColor(75,205,62)

    log.info("Filling up RCA Details.")
    #RCA Details
    h1 = doc.add_heading('Root Cause Analysis Details', 1)
    h1.style.font.color.rgb = RGBColor(0,151,117)
    h2 = doc.add_heading('Root Cause', 2)
    h2.style.font.color.rgb = RGBColor(1,91,126)
    rcaparagraph1 = doc.add_paragraph(rca_det[0])
    h2 = doc.add_heading('RCA Executive Summary', 2)
    h2.style.font.color.rgb = RGBColor(1,91,126)
    rcaparagraph1 = doc.add_paragraph(rca_det[1])
    h2 = doc.add_heading('Investigation & Resolution', 2)
    h2.style.font.color.rgb = RGBColor(1,91,126)
    rcaparagraph1 = doc.add_paragraph(rca_det[2])
    h2 = doc.add_heading('Contributing Factors', 2)
    h2.style.font.color.rgb = RGBColor(1,91,126)
    rcaparagraph1 = doc.add_paragraph(rca_det[3])

    log.info("Filling up Action Items.")
    #Filling up Action Items
    h1 = doc.add_heading("Action Items", 1)
    h1.style.font.color.rgb = RGBColor(0,151,117)
    action_item_table = doc.add_table(rows=1, cols=5)
    action_item_table.style = "Table Grid"
    columns = action_item_table.rows[0].cells
    columns[0].text = "Actions"
    columns[1].text = "Description"
    columns[2].text = "Owner"
    columns[3].text = "Date"
    columns[4].text = "Status"

    for x in action_items:
        row = action_item_table.add_row().cells
        row[0].text = x["Actions"]
        row[1].text = x["Description"]
        row[2].text = x["Owner"]
        row[3].text = x["Date"]
        row[4].text = x["Status"]

    log.info("5 Whys.")
    #Filling up Action Items
    h1 = doc.add_heading("5 Whys", 1)
    h1.style.font.color.rgb = RGBColor(0,151,117)
    whyparagraph = doc.add_paragraph(rca_why)


    log.info("Filling up incident timeline.")
    #incident timeline
    h1 = doc.add_heading('Incident Timeline', 1)
    h1.style.font.color.rgb = RGBColor(0,151,117)
    timeline_table = doc.add_table(rows=1, cols=3)
    timeline_table.style = "Table Grid"
    columns = timeline_table.rows[0].cells
    columns[0].text = "Date"
    columns[1].text = "Time"
    columns[2].text = "Contents"
    
    log.info(inc_timeline[0])

    for x in inc_timeline:
        log.info(f"processing {x}")
        row = timeline_table.add_row().cells
        row[0].text = x["Date"]
        row[1].text = x["Time"]
        row[2].text = x["Contents"]

    # doc.font.color.rgb = RGBColor(75,205,62)
    doc.add_page_break()
    doc.save('output.docx')

def convert_word_to_pdf(filepath, filename='output'):
    '''
    convert word to pdf

    parameters:
        filepath - docx filepath
    
    return:
        path - pdf path
    '''
    log.info("Converting DOCX to PDF")
    path = f'{filename}.pdf'

    log.info(f"pdf path and name: {path}")
    convert(filepath, path)

    return path

def convert_word_to_pdf_unix(filepath, filename='output'):
    '''
    convert word to pdf

    parameters:
        filepath - docx filepath
    
    return:
        path - pdf path
    '''

    path = f"{filename}.pdf"
    log.info("Converting DOCX to PDF")
    docfile = aw.Document(filepath)
    docfile.save(path)

    log.info(f"pdf path and name: {path}")

    return path